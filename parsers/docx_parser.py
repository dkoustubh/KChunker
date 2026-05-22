from pathlib import Path
from typing import Any

from plugin_system.base import BaseParser, ExtractedDocument
from utils.logging import logger

docx: Any = None
try:
    import docx as docx_module

    docx = docx_module
except ImportError:
    pass


class DOCXParser(BaseParser):
    """Parses Word documents (DOCX), extracting paragraphs and tables."""

    @property
    def name(self) -> str:
        return "docx_parser"

    @property
    def version(self) -> str:
        return "1.0.0"

    def can_parse(self, file_path: Path, mime_type: str) -> bool:
        suffix = file_path.suffix.lower()
        return suffix == ".docx" or "word" in mime_type

    async def parse(self, file_path: Path) -> ExtractedDocument:
        logger.info("Parsing DOCX document", path=str(file_path))

        if docx is None:
            raise ImportError(
                "python-docx is required for parsing Word files. "
                "Please run `uv sync` to install dependencies."
            )

        raw_text_parts: list[str] = []
        extracted_tables: list[dict[str, Any]] = []

        try:
            doc = docx.Document(str(file_path))

            # 1. Parse Paragraphs
            for p in doc.paragraphs:
                text = p.text.strip()
                if text:
                    raw_text_parts.append(text)

            # 2. Parse Tables
            for table_idx, table in enumerate(doc.tables):
                table_data: list[list[str]] = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)

                extracted_tables.append({"table_index": table_idx, "data": table_data})

                # Append table text representation to raw text for semantic search
                table_str = "\n".join([" | ".join(row) for row in table_data])
                raw_text_parts.append(f"\n[Table {table_idx}]\n{table_str}\n")

        except Exception as e:
            logger.error("Failed to parse DOCX file", path=str(file_path), error=str(e))
            raise e

        combined_text = "\n\n".join(raw_text_parts)

        return ExtractedDocument(
            document_name=file_path.name,
            document_path=str(file_path),
            document_type="DOCX",
            raw_text=combined_text,
            pages=[{"page_num": 1, "text": combined_text}],
            tables=extracted_tables,
            emails=[],
            metadata={
                "num_paragraphs": len(doc.paragraphs),
                "num_tables": len(extracted_tables),
            },
        )
